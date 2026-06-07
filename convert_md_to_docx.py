#!/usr/bin/env python
"""Convert TECHNICAL_MENTOR_MASTERCLASS.md to .docx format"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

md_file = r'C:\papper2code\TECHNICAL_MENTOR_MASTERCLASS.md'
docx_file = r'C:\papper2code\TECHNICAL_MENTOR_MASTERCLASS.docx'

print(f"Reading: {md_file}")
with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

print(f"Creating Word document...")
doc = Document()

# Set default styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Split by lines
lines = content.split('\n')

in_code_block = False
code_lines = []
current_part = 0

for i, line in enumerate(lines):
    # Track code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            # End code block
            if code_lines:
                code_text = '\n'.join(code_lines).strip()
                p = doc.add_paragraph(code_text, style='Normal')
                p.paragraph_format.left_indent = Inches(0.3)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        continue
    
    if in_code_block:
        code_lines.append(line)
        continue
    
    # Skip empty lines
    if not line.strip():
        doc.add_paragraph()
        continue
    
    # Detect parts for page breaks
    if line.startswith('# PART '):
        current_part += 1
        if current_part > 1:
            doc.add_page_break()
        h = doc.add_heading(line[2:].strip(), level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        continue
    
    # Headings
    if line.startswith('# '):
        h = doc.add_heading(line[2:].strip(), level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=4)
    elif line.startswith('##### '):
        doc.add_heading(line[6:].strip(), level=5)
    
    # Bullet points
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:].strip()
        if text:
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
    
    # Numbered lists
    elif re.match(r'^\d+\.\s', line.strip()):
        match = re.match(r'^(\d+)\.\s(.+)', line.strip())
        if match:
            text = match.group(2)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            doc.add_paragraph(text, style='List Number')
    
    # Regular paragraphs
    else:
        text = line.strip()
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'~~(.+?)~~', r'\1', text)
        # Remove control characters and null bytes
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t\r')
        if text and text not in ['---', '---\n']:
            try:
                doc.add_paragraph(text)
            except ValueError:
                # Skip problematic lines
                pass

doc.save(docx_file)
file_size = os.path.getsize(docx_file) / 1024
print(f"\n✅ SUCCESS!")
print(f"📄 File: {os.path.basename(docx_file)}")
print(f"📊 Size: {file_size:.1f} KB")
print(f"📍 Location: {docx_file}")
